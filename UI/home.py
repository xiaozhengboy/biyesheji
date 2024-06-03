# 导入必要的模块
from datetime import datetime

from PyQt5.QtCore import QTimer
from PyQt5.QtGui import *
from PyQt5.QtWidgets import *
import sys

import efficientnet.tongue_shese.predict as sz
import efficientnet.tongue_taise.predict as ts
import efficientnet.tongue_chihen.predict as ch
import efficientnet.tongue_liewen.predict as lw

import os
from UI.zhuye import Ui_HomePageWindow
from UI.xiaozheng import Ui_MainWindow
from docx import Document
from docx.shared import Inches
from docx.enum.text import WD_PARAGRAPH_ALIGNMENT
from docxtpl import DocxTemplate, InlineImage
from docx.shared import Mm
import cv2
from PIL import Image
# 允许重复加载库
os.environ["KMP_DUPLICATE_LIB_OK"]="TRUE"

class mainWin(QMainWindow, Ui_HomePageWindow):

    def __init__(self,parent=None):

        super(mainWin, self).__init__(parent)
        self.setupUi(self)
        self.pushButton.clicked.connect(self.login)

    def login(self):

        main_win.close()
        second_main.show()

class secondmain(QMainWindow, Ui_MainWindow):

    def __init__(self, parent=None):

        super(secondmain,self).__init__(parent)
        self.setupUi(self)
        # 上传舌象
        self.shangchuan.clicked.connect(self.openimage)
        self.shangchuan.setFlat(True)

        # 点击拍照
        self.paizhao.clicked.connect(self.photo)

        # 下载
        self.xiazai.clicked.connect(self.download)
        self.xiazai.setFlat(True)

        # 检测
        self.start.clicked.connect(self.start_predict)

    def center_crop(self,img, target_size):
        width, height = img.size
        target_width, target_height = target_size
        if width < target_width or height < target_height:
            return img
        left = (width - target_width) // 2
        top = (height - target_height) // 2
        right = (width + target_width) // 2
        bottom = (height + target_height) // 2
        return img.crop((left, top, right, bottom))

    def photo(self):
        self.sz_color.setText("")
        self.ts_color.setText("")
        self.liewen.setText("")
        self.chihen.setText("")
        self.wb_textarea.setText("")
        self.xiazai.setText("打印报告")

        # cap = cv2.VideoCapture(0)
        # while True:
        #     ret, frame = cap.read()
        #     if ret:
        #         rgb_image = cv2.cvtColor(frame, cv2.COLOR_BGR2RGB)
        #         h, w, ch = rgb_image.shape
        #         bytes_per_line = ch * w
        #         qt_image = QImage(rgb_image.data, w, h, bytes_per_line, QImage.Format_RGB888)
        #         pixmap = QPixmap.fromImage(qt_image)
        #         self.uesr_tongue.setPixmap(pixmap)
        #
        #         # cv2.imshow('拍照中，请伸出舌头。按Q键退出', frame)
        #         if cv2.waitKey(1) & 0xFF == ord('q'):
        #             cv2.imwrite('user_load.jpg', frame)
        #             break
        # cap.release()
        # cv2.destroyAllWindows()

        cap = cv2.VideoCapture(0)
        self.wb_textarea.setText('请将舌头放入绿色框中，按ENTER键即可完成拍照。')
        while True:
            ret, frame = cap.read()
            height, width, _ = frame.shape

            box_width, box_height = 400, 400
            start_x = (width - box_width) // 2
            start_y = (height - box_height) // 2
            end_x = start_x + box_width
            end_y = start_y + box_height
            cv2.rectangle(frame, (start_x, start_y), (end_x, end_y), (0, 255, 0), 2)

            cv2.imshow('Camera', frame)
            # if cv2.waitKey(1) & 0xFF == ord('q'):
            if cv2.waitKey(1) & 0xFF == 13:
                cv2.imwrite('user_load.jpg', frame)
                img = Image.open("user_load.jpg")
                cropped_img = self.center_crop(img, (396, 396))
                qimage = QImage(cropped_img.size[0], cropped_img.size[1], QImage.Format_RGB888)
                for x in range(cropped_img.size[0]):
                    for y in range(cropped_img.size[1]):
                        r, g, b = cropped_img.getpixel((x, y))
                        qimage.setPixel(x, y, qRgb(r, g, b))
                # cropped_image = frame[start_y:end_y, start_x:end_x]
                jpg = QPixmap.fromImage(qimage)
                jpg.save('./user_load.jpg')
                self.wb_textarea.setText("")
                self.uesr_tongue.setPixmap(jpg)
                break
            if cv2.getWindowProperty("Camera", cv2.WND_PROP_VISIBLE) < 1:
                break
        cap.release()
        cv2.destroyAllWindows()


    def openimage(self):
        self.sz_color.setText("")
        self.ts_color.setText("")
        self.liewen.setText("")
        self.chihen.setText("")
        self.wb_textarea.setText("")
        self.xiazai.setText("打印报告")
        # 打开文件对话框并获取选择的文件名和类型
        imgName,imgType = QFileDialog.getOpenFileName(self,"打开图片", "../data/test", "*.jpg;*.tif;*.png;;All Files(*)")
        if imgName == "":
            return 0
        # 加载并调整图像大小，保存并显示在界面上
        jpg = QPixmap(imgName).scaled(self.uesr_tongue.width(),self.uesr_tongue.height())
        jpg.save('./user_load.jpg')
        self.wb_textarea.setText("")
        self.uesr_tongue.setPixmap(jpg)


    def start_predict(self):
        self.sz_color.setText("")
        self.ts_color.setText("")
        self.liewen.setText("")
        self.chihen.setText("")
        self.xiazai.setText("打印报告")


        # 调用模型进行预测并获取结果
        lw_class, lw_prob = lw.main('./user_load.jpg')  # 裂纹
        ch_class, ch_prob = ch.main('./user_load.jpg')  # 齿痕
        sz_class, sz_prob = sz.main('./user_load.jpg')  # 舌质
        ts_class, ts_prob = ts.main('./user_load.jpg') # 苔色

        print(lw_class, lw_prob)
        print(ch_class,ch_prob)
        print(sz_class, sz_prob)
        print(ts_class, ts_prob)


        # 根据预测结果更新界面显示
        if lw_class == 'you':
            lw_res = "有裂纹"
        else:
            lw_res = "无裂纹"


        if ch_class == 'you_kuo':
            ch_res = "有齿痕"
        else:
            ch_res = "无齿痕"


        if ts_class == "bobaitai":
            ts_res = "薄白苔"
        elif ts_class == "huangtai":
            ts_res = "黄苔"
        elif ts_class == "houbaitai":
            ts_res = "厚白苔"
        elif ts_class == "huiheitai":
            ts_res = "灰黑苔"
        else:
            ts_res = "无苔"


        if sz_class == "danbai":
            sz_res = "淡白舌"
        elif sz_class == "danhong":
            sz_res = "淡红舌"
        elif sz_class == "jiang":
            sz_res = "绛舌"
        elif sz_class == "qingzi":
            sz_res = "青紫舌"
        else:
            sz_res = "鲜红舌"


        self.liewen.setText("{} {:.3}".format(lw_res,lw_prob))
        self.chihen.setText("{} {:.3}".format(ch_res,ch_prob))
        self.sz_color.setText("{} {:.3}".format(sz_res,sz_prob))
        self.ts_color.setText("{} {:.3}".format(ts_res,ts_prob))


        # self.liewen.setText("{}".format(crack_res,crack_prob))
        # self.chihen.setText("{}".format(indent_res,coated_prob))
        # self.sz_color.setText("{}".format(color_res,color_prob))
        # self.ts_color.setText("{}".format(coated_res,indent_prob))


        # 根据预测结果提供建议文本
        # 淡红舌
        if sz_class == 'danhong' and ts_class == 'bobaitai':
            self.wb_textarea.setText("淡红舌，薄白苔:\n健康人；风寒表征；病势轻浅")
        elif sz_class == 'danhong' and ts_class == 'houbaitai':
            self.wb_textarea.setText("淡红舌，厚白苔:\n湿浊痰饮内停；食积胃肠；寒湿痹症")
        elif sz_class == 'danhong' and ts_class == 'huangtai' :
            self.wb_textarea.setText("淡红舌，黄苔:\n里热轻证;风热表证")
        elif sz_class == 'danhong' and ts_class == 'huiheitai' :
            self.wb_textarea.setText("淡红舌，灰黑苔:\n寒证;阳虚")
        elif sz_class == 'danhong' and ts_class == 'wutai' :
            self.wb_textarea.setText("淡红舌，无苔:\n")

        # 淡白舌
        if sz_class == 'danbai' and ts_class == 'bobaitai':
            self.wb_textarea.setText("淡白舌，薄白苔:\n阳气不足;气血虚弱")
        elif sz_class == 'danbai' and ts_class == 'houbaitai':
            self.wb_textarea.setText("淡白舌，厚白苔:\n脾胃虚弱,痰湿停聚")
        elif sz_class == 'danbai' and ts_class == 'huangtai' :
            self.wb_textarea.setText("淡白舌，黄苔:\n")
        elif sz_class == 'danbai' and ts_class == 'huiheitai' :
            self.wb_textarea.setText("淡白舌，灰黑苔:\n阳虚内寒;痰湿内停")
        elif sz_class == 'danbai' and ts_class == 'wutai' :
            self.wb_textarea.setText("淡白舌，无苔:\n久病阳衰;气血俱虚;脾胃虚寒")

        # 绛舌
        if sz_class == 'jiang' and ts_class == 'bobaitai':
            self.wb_textarea.setText("绛舌，薄白苔:\n")
        elif sz_class == 'jiang' and ts_class == 'houbaitai':
            self.wb_textarea.setText("绛舌，厚白苔:\n")
        elif sz_class == 'jiang' and ts_class == 'huangtai':
            self.wb_textarea.setText("绛舌，黄苔:\n邪热深重;胃肠热结")
        elif sz_class == 'jiang' and ts_class == 'huiheitai':
            self.wb_textarea.setText("绛舌，灰黑苔:\n热极伤阴")
        elif sz_class == 'jiang' and ts_class == 'wutai':
            self.wb_textarea.setText("绛舌，无苔:\n热入血分;阴虚火旺")

        # 青紫舌
        if sz_class == 'qingzi' and ts_class == 'bobaitai':
            self.wb_textarea.setText("青紫舌，薄白苔:\n阳衰寒盛;气血凝滞")
        elif sz_class == 'qingzi' and ts_class == 'houbaitai':
            self.wb_textarea.setText("青紫舌，厚白苔:\n")
        elif sz_class == 'qingzi' and ts_class == 'huangtai':
            self.wb_textarea.setText("青紫舌，黄苔:\n热极津枯")
        elif sz_class == 'qingzi' and ts_class == 'huiheitai':
            self.wb_textarea.setText("青紫舌，灰黑苔:\n热毒深重,津液大伤")
        elif sz_class == 'qingzi' and ts_class == 'wutai':
            self.wb_textarea.setText("青紫舌，无苔:\n")

        # 鲜红舌
        if sz_class == 'xianhong' and ts_class == 'bobaitai':
            self.wb_textarea.setText("鲜红舌，薄白苔:\n")
        elif sz_class == 'xianhong' and ts_class == 'houbaitai':
            self.wb_textarea.setText("鲜红舌，厚白苔:\n正气亏虚;湿热未净&里热夹痰湿;阴虚兼痰湿")
        elif sz_class == 'xianhong' and ts_class == 'huangtai':
            self.wb_textarea.setText("鲜红舌，黄苔:\n里热证,津液已伤&气分热盛,阴液耗损&湿热内蕴;痰热互结")
        elif sz_class == 'xianhong' and ts_class == 'huiheitai':
            self.wb_textarea.setText("鲜红舌，灰黑苔:\n津枯血燥")
        elif sz_class == 'xianhong' and ts_class == 'wutai':
            self.wb_textarea.setText("鲜红舌，无苔:\n")

    def download(self):
        text = self.wb_textarea.toPlainText()
        images = './user_load.jpg'  # 保存在本地的图片
        if text != '':
            try:
                self.xiazai.setText('报告生成中')
                tpl = DocxTemplate('muban.docx')
                now = datetime.now()
                now_time = now.strftime('%Y-%m-%d %H:%M:%S')
                now_filename = now.strftime('%Y%m%d%H%M%S')
                # 设置好各标签需要填写的内容
                context = {
                    'date': now_time,
                    'taise': self.ts_color.text(),
                    'shese': self.sz_color.text(),
                    'chihen': self.chihen.text(),
                    'liewen': self.liewen.text(),
                    'img': InlineImage(tpl, images, width=Mm(60)),
                    'message': text
                }
                # 将标签内容填入模板中
                tpl.render(context)
                # 保存
                tpl.save(rf'./result/{now_filename}.docx')
                self.xiazai.setText('打印成功')
            except Exception as e:
                print(e)
                self.xiazai.setText('打印失败')
        else:
            self.xiazai.setText('请先检测')
            self.wb_textarea.setText("请先完成检测")

    # def download(self):
    #     text = self.wb_textarea.toPlainText()
    #     images = './user_load.jpg'  # 保存在本地的图片
    #     if text != '':
    #         try:
    #             self.xiazai.setText('下载中')
    #             doc = Document()  # doc对象
    #             paragraph = doc.add_paragraph()
    #             paragraph.alignment = WD_PARAGRAPH_ALIGNMENT.CENTER
    #             run = paragraph.add_run("")
    #             run.add_picture(images, width=Inches(2))
    #             doc.add_paragraph(text)  # 添加文字
    #             doc.save(rf'./result/{images}.docx')  # 保存路径
    #             self.xiazai.setText('下载成功')
    #         except Exception as e:
    #             self.xiazai.setText('下载失败')
    #     else:
    #         self.xiazai.setText('请先检测')



if __name__ == '__main__':
    app = QApplication(sys.argv)
    main_win = mainWin()
    second_main = secondmain()
    main_win.show()
    sys.exit(app.exec_())